from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Reporte de Verificaci\u00f3n y Comprobaci\u00f3n del Proyecto BITSAC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'Fecha: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}')
doc.add_paragraph('=' * 80)
doc.add_paragraph('')

# 1. INFORMACION GENERAL
doc.add_heading('1. Informaci\u00f3n General del Proyecto', 1)
doc.add_paragraph('Nombre del proyecto: BITSAC (SIGO-GCAM)')
doc.add_paragraph('Descripci\u00f3n: Sistema Inteligente de Gesti\u00f3n Operativa para el Cuerpo de Agentes de Control Municipal de Guayaquil')
doc.add_paragraph('')

info_table = doc.add_table(rows=10, cols=2)
info_table.style = 'Light Grid Accent 1'
info_table.cell(0,0).text = 'Componente'
info_table.cell(0,1).text = 'Detalle'
info_table.cell(1,0).text = 'Backend'
info_table.cell(1,1).text = 'Node.js (Express 5.2.1) - JavaScript (CommonJS)'
info_table.cell(2,0).text = 'Base de Datos'
info_table.cell(2,1).text = 'SQL Server (ODBC Driver 18)'
info_table.cell(3,0).text = 'Mobile'
info_table.cell(3,1).text = 'Flutter 3.44.4 / Dart 3.12.2'
info_table.cell(4,0).text = 'Sistemas Operativos'
info_table.cell(4,1).text = 'Android, iOS, Web, Windows, macOS, Linux'
info_table.cell(5,0).text = 'Arquitectura Backend'
info_table.cell(5,1).text = '4 Capas: Routes \u2192 Controllers \u2192 Services \u2192 Repositories'
info_table.cell(6,0).text = 'Arquitectura Mobile'
info_table.cell(6,1).text = 'Feature-first + Clean Architecture simplificada'
info_table.cell(7,0).text = 'Autenticaci\u00f3n'
info_table.cell(7,1).text = 'JWT + bcrypt + RBAC (Role-Based Access Control)'
info_table.cell(8,0).text = 'M\u00f3dulos'
info_table.cell(8,1).text = 'Personal, Eventos, Anuncios, Cartillas, Insignias, Administraci\u00f3n, Cat\u00e1logos'
info_table.cell(9,0).text = 'L\u00edneas de c\u00f3digo'
info_table.cell(9,1).text = 'Backend: ~3,590 l\u00edneas | Mobile: ~5,000+ l\u00edneas | SQL: ~1,511 l\u00edneas'

doc.add_paragraph('')

# 2. LO QUE FUNCIONA
doc.add_heading('2. Lo que Funciona Correctamente', 1)

works = [
    ('2.1 Backend - Estructura y Arquitectura', [
        'Arquitectura en 4 capas bien definida (Routes \u2192 Controllers \u2192 Services \u2192 Repositories)',
        'Separaci\u00f3n de responsabilidades correcta y consistente en todos los m\u00f3dulos',
        'Uso de Express 5 (versi\u00f3n m\u00e1s reciente) con middlewares globales adecuados',
        'Manejo de errores global con c\u00f3digos HTTP apropiados (400, 401, 403, 404, 413, 500)',
        'Sistema de auditor\u00eda funcional que registra acciones CRUD autom\u00e1ticamente',
        'Middleware de autenticaci\u00f3n JWT con verificaci\u00f3n de token y expiraci\u00f3n',
        'Sistema RBAC con 3 funciones: requireAuth, requirePermission, requireAnyPermission'
    ]),
    ('2.2 Backend - M\u00f3dulos Funcionales', [
        'M\u00f3dulo de Autenticaci\u00f3n: login con JWT (8h expiraci\u00f3n), cambio de contrase\u00f1a, hash bcrypt',
        'M\u00f3dulo de Personal: CRUD completo, b\u00fasqueda, perfiles, cambio de estado, reset de contrase\u00f1a',
        'M\u00f3dulo de Eventos: CRUD, cambio de estado, convocatoria de personal, vista de confirmados',
        'M\u00f3dulo de Anuncios: Publicaci\u00f3n con asignaci\u00f3n personal, expiraci\u00f3n, notificaciones',
        'M\u00f3dulo de Cartillas: Generaci\u00f3n de reportes de novedades con formato SAC',
        'M\u00f3dulo de Insignias: Sistema de gamificaci\u00f3n con 15 insignias escalonadas',
        'M\u00f3dulo de Administraci\u00f3n: CRUD de cat\u00e1logos, roles, permisos, lugares, EAS, m\u00f3viles, asignaciones',
        'Vista de alertas de mantenimiento preventivo de m\u00f3viles'
    ]),
    ('2.3 Backend - Base de Datos', [
        'Migraciones SQL con control de versiones (formato YYYYMMDD)',
        'Scripts idempotentes (usan IF ... COL_LENGTH para evitar errores en re-ejecuci\u00f3n)',
        'Uso de vistas SQL Server para abstraer consultas complejas (vw_personal_detalle, vw_personal_operativo, etc.)',
        'Sistema RBAC completo con ~80 permisos en 24 m\u00f3dulos',
        'Transacciones SQL en operaciones cr\u00edticas (crear/actualizar eventos con personal asociado)',
        'Integridad referencial con claves for\u00e1neas y \u00edndices \u00fanicos filtrados'
    ]),
    ('2.4 Mobile - Estructura y UI', [
        'Aplicaci\u00f3n multi-plataforma (Android, iOS, Web, Windows, macOS, Linux)',
        'UI adaptativa responsiva (Web vs Mobile seg\u00fan ancho de pantalla)',
        'Pantalla de Login con validaci\u00f3n de campos vac\u00edos y ocultaci\u00f3n de teclado',
        'Dashboard con men\u00fa lateral y m\u00f3dulos funcionales (Eventos, Cartillas, Insignias, Administraci\u00f3n)',
        'Navegaci\u00f3n con gesti\u00f3n de estado usando setState y mounted checks',
        'Manejo de errores con timeout y mensajes descriptivos al usuario',
        'Soporte nativo para web con implementaciones condicionales (stub vs web)',
        'Tema institucional consistente con colores SEGURA EP'
    ]),
    ('2.5 Scripts y Herramientas', [
        'Script de prueba de conexi\u00f3n ODBC con 8 variantes de conexi\u00f3n',
        'Script de auditor\u00eda de esquema SQL vs c\u00f3digo fuente con fallback est\u00e1tico',
        'Parches para correcci\u00f3n de validaci\u00f3n en login (auth_scr.dart)'
    ]),
]

for section_title, items in works:
    doc.add_heading(section_title, 2)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run('\u2705 ' + item)
        run.font.color.rgb = RGBColor(0, 128, 0)

# 3. ERRORES DETECTADOS
doc.add_heading('3. Errores Detectados', 1)

errors = [
    ('3.1 CR\u00cdTICOS', [
        ('[CRIT-01] Contrase\u00f1a por defecto insegura', 'media', 
         'El script reset_usuarios_seed.sql usa la c\u00e9dula como contrase\u00f1a inicial. Adem\u00e1s, si la columna password_hash no existe, el login acepta la c\u00e9dula como contrase\u00f1a v\u00e1lida (validarClave, l\u00ednea 275-285 de auth.routes.js).',
         'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:275 | C:\\\\database\\\\*_reset_usuarios_seed.sql'),
        ('[CRIT-02] JWT_SECRET hardcodeado y expuesto', 'media',
         'El JWT_SECRET en .env es "sigo_bitsac_creado_Lunatics" y el fallback es "sigo_gcam_secret". El .env con datos reales est\u00e1 en el repositorio.',
         'C:\\\\backend\\\\.env | C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:97'),
        ('[CRIT-03] Sin limitaci\u00f3n de intentos de login', 'alta',
         'No hay rate limiting ni bloqueo por intentos fallidos en el endpoint /api/auth/login. Vulnerable a ataques de fuerza bruta.',
         'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:15'),
        ('[CRIT-04] Inyecci\u00f3n SQL potencial en admin.repository.js', 'media',
         'La funci\u00f3n insertarBasico y cambiarActivo usan interpolaci\u00f3n directa del nombre de tabla: INSERT INTO ${tabla}...',
         'C:\\\\backend\\\\src\\\\repositories\\\\admin.repository.js:459-480'),
        ('[CRIT-05] Token JWT sin renovaci\u00f3n ni refresh token', 'media',
         'El token expira en 8h sin mecanismo de refresh. Al expirar, el usuario pierde la sesi\u00f3n forzosamente.',
         'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:95-99'),
    ]),
    ('3.2 GRAVES', [
        ('[ERR-01] Conexi\u00f3n ODBC sin pool', 'alta',
         'Cada operaci\u00f3n de base de datos abre y cierra una conexi\u00f3n individual. Sin pool de conexiones, el rendimiento se degrada severamente con m\u00faltiples usuarios.',
         'Todos los repositories'),
        ('[ERR-02] Error HTTP 500 en lugar de 404 en obtenerPerfil', 'media',
         'El controller de personal usa status(404) pero si el service lanza otro error (ej. BD ca\u00edda), responde 500, lo cual es correcto, pero el mensaje de error expone detalles internos al cliente.',
         'C:\\\\backend\\\\src\\\\controllers\\\\personal.controller.js:84-97'),
        ('[ERR-03] Falta sanitizaci\u00f3n de entrada en b\u00fasqueda', 'baja',
         'El endpoint de b\u00fasqueda usa LIKE con el texto directamente. Aunque el driver ODBC escapa par\u00e1metros, no hay validaci\u00f3n de longitud m\u00ednima/m\u00e1xima.',
         'C:\\\\backend\\\\src\\\\repositories\\\\personal.repository.js:77-97'),
        ('[ERR-04] Sin validaci\u00f3n de permisos en cambio de contrase\u00f1a', 'media',
         'El endpoint /api/auth/change-password solo requiere autenticaci\u00f3n (requireAuth) pero no verifica permisos adicionales.',
         'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:120'),
        ('[ERR-05] Variables de entorno sin validaci\u00f3n exhaustiva', 'media',
         'db.js valida DB_SERVER y DB_DATABASE, pero no valida DB_DRIVER, DB_ENCRYPT ni DB_CONNECTION_TIMEOUT.',
         'C:\\\\backend\\\\src\\\\config\\\\db.js'),
    ]),
    ('3.3 MOBILE (Flutter/Dart)', [
        ('[MOB-01] Gesti\u00f3n de estado d\u00e9bil (solo setState)', 'alta',
         'Toda la app usa setState para manejo de estado. Sin Provider, BLoC, Riverpod ni ning\u00fan gestor de estado. Esto causa renderizados innecesarios y dificulta el mantenimiento.',
         'Todos los StatefulWidget'),
        ('[MOB-02] AuthSession con singleton est\u00e1tico puro', 'media',
         'AuthSession es una clase con campos static. El token se pierde al cerrar la app. No persiste en SharedPreferences ni SecureStorage.',
         'C:\\\\mobile\\\\lib\\\\core\\\\auth\\\\auth_session.dart'),
        ('[MOB-03] Sin manejo de errores HTTP granulares', 'media',
         'Todos los errores HTTP se tratan como Exception gen\u00e9rica. No hay distinci\u00f3n entre 401 (no autorizado), 403 (prohibido), 404, 500, etc.',
         'C:\\\\mobile\\\\lib\\\\core\\\\api\\\\api_client.dart:138-169'),
        ('[MOB-04] Router vac\u00edo (app_rtr.dart sin implementaci\u00f3n)', 'alta',
         'El archivo app_rtr.dart est\u00e1 completamente vac\u00edo. La navegaci\u00f3n usa Navigator.pushReplacement imperativo en lugar de un router declarativo.',
         'C:\\\\mobile\\\\lib\\\\core\\\\rtr\\\\app_rtr.dart'),
        ('[MOB-05] Sin autologout al recibir 401', 'alta',
         'Si el backend responde 401 (token expirado), la app no redirige autom\u00e1ticamente al login. El usuario ve un error gen\u00e9rico.',
         'C:\\\\mobile\\\\lib\\\\core\\\\api\\\\api_client.dart'),
        ('[MOB-06] NotifReadStore usa Set en memoria (se pierde al cerrar app)', 'media',
         'Las notificaciones le\u00eddas solo se marcan en un Set est\u00e1tico en memoria. Al cerrar la app, se pierde el estado de le\u00eddo.',
         'C:\\\\mobile\\\\lib\\\\core\\\\notif\\\\notif_read_store.dart'),
    ]),
    ('3.4 BASE DE DATOS', [
        ('[DB-01] Reset_usuarios_seed.sql elimina TODOS los datos sin advertencia', 'alta',
         'El script elimina datos existentes sin preguntar ni hacer backup. Peligroso en producci\u00f3n.',
         'C:\\\\database\\\\*_reset_usuarios_seed.sql'),
        ('[DB-02] Sin \u00edndices expl\u00edcitos en tablas cr\u00edticas', 'media',
         'No se definen \u00edndices en columnas usadas frecuentemente en WHERE/JOIN (correo_institucional, cedula, evento_presidente.evento_id, etc.)',
         'Archivos SQL en C:\\\\database'),
        ('[DB-03] Contracci\u00f3n de nombres inconsistente', 'baja',
         'Mezcla de estilos: fecha_creacion vs fecha_actualizacion vs createdAt. Algunas tablas tienen ambas columnas, otras solo una.',
         'Archivos SQL en C:\\\\database'),
    ]),
]

for section_title, items in errors:
    doc.add_heading(section_title, 2)
    for title, severity, desc, location in items:
        p = doc.add_paragraph()
        run = p.add_run(f'\u274C {title}')
        run.bold = True
        run.font.color.rgb = RGBColor(200, 0, 0)
        
        p2 = doc.add_paragraph()
        severity_colors = {'alta': RGBColor(255, 0, 0), 'media': RGBColor(255, 140, 0), 'baja': RGBColor(200, 180, 0)}
        run2 = p2.add_run(f'[Severidad: {severity.upper()}] ')
        run2.font.color.rgb = severity_colors.get(severity, RGBColor(0,0,0))
        run2.bold = True
        
        doc.add_paragraph(desc)
        doc.add_paragraph(f'Ubicaci\u00f3n: {location}')

# 4. LO OBSOLETO
doc.add_heading('4. C\u00f3digo Obsoleto / Deprecated', 1)

obs = [
    ('[OBS-01] Dependencias duplicadas de SQL Server', 'media',
     'package.json incluye odbc (driver principal), msnodesqlv8 (driver nativo Windows) y mssql (driver alternativo). Solo se usa odbc realmente. Los otros dos son redundantes.',
     'C:\\\\backend\\\\package.json'),
    ('[OBS-02] Scripts con rutas relativas a node_modules', 'baja',
     'test_odbc_connection.js require("../backend/node_modules/odbc") en lugar de usar la resoluci\u00f3n normal de m\u00f3dulos.',
     'C:\\\\scripts\\\\test_odbc_connection.js:1'),
    ('[OBS-03] Parche duplicado (auth_01_correct.patch vs auth_01.patch)', 'baja',
     'Hay dos parches casi id\u00e9nticos. auth_01_correct.patch es una versi\u00f3n incompleta. Deber\u00eda eliminarse.',
     'C:\\\\patch'),
    ('[OBS-04] .agents/ directorio vac\u00edo', 'baja',
     'El directorio .agents/ est\u00e1 vac\u00edo. Si no se va a usar, deber\u00eda eliminarse.',
     'C:\\\\.agents'),
    ('[OBS-05] estadisticas.txt vac\u00edo en backend/', 'baja',
     'Archivo estadisticas.txt en la ra\u00edz del backend est\u00e1 vac\u00edo (0 bytes). Sin uso aparente.',
     'C:\\\\backend\\\\estadisticas.txt'),
    ('[OBS-06] README.md vac\u00edo', 'baja',
     'El README.md del proyecto ra\u00edz est\u00e1 completamente vac\u00edo.',
     'C:\\\\README.md'),
    ('[OBS-07] assets/anim/ y assets/icn/ vac\u00edos', 'baja',
     'Las carpetas assets/anim/ y assets/icn/ en el proyecto mobile est\u00e1n vac\u00edas pero referenciadas en pubspec.yaml.',
     'C:\\\\mobile\\\\assets\\\\anim | C:\\\\mobile\\\\assets\\\\icn | pubspec.yaml:60-62'),
]

for title, severity, desc, location in obs:
    p = doc.add_paragraph()
    run = p.add_run(f'\u26A0\uFE0F {title}')
    run.bold = True
    run.font.color.rgb = RGBColor(180, 120, 0)
    doc.add_paragraph(f'[Severidad: {severity.upper()}] {desc}')
    doc.add_paragraph(f'Ubicaci\u00f3n: {location}')

# 5. REDUNDANCIA
doc.add_heading('5. C\u00f3digo Redundante', 1)

red = [
    ('[RED-01] Ruta /api/probar-db duplicada', 'baja',
     'La ruta /probar-db se registra dos veces: como /probar-db y como /api/probar-db en index.js.',
     'C:\\\\backend\\\\index.js:75-76'),
    ('[RED-02] Validaci\u00f3n de c\u00e9dula y correo duplicada en personal.service.js', 'media',
     'La validaci\u00f3n de c\u00e9dula (10 d\u00edgitos) y correo (regex) est\u00e1 repetida en las funciones actualizarPerfil, crear y actualizar.',
     'C:\\\\backend\\\\src\\\\services\\\\personal.service.js:48-55, 106-113, 150-157'),
    ('[RED-03] Contrase\u00f1a por defecto duplicada en auth.routes y personal.service', 'media',
     'La funci\u00f3n passwordInicial/passwordDesdeFecha est\u00e1 implementada dos veces con la misma l\u00f3gica.',
     'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:450-459 | C:\\\\backend\\\\src\\\\services\\\\personal.service.js:218-233'),
    ('[RED-04] Funci\u00f3n tieneColumnaFoto duplicada', 'media',
     'La funci\u00f3n para verificar existencia de columna foto_perfil_url est\u00e1 implementada en auth.routes.js y personal.repository.js.',
     'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:461-467 | C:\\\\backend\\\\src\\\\repositories\\\\personal.repository.js:196-202'),
    ('[RED-05] Permisos hardcodeados en auth.routes.js que duplican los de la BD', 'alta',
     'La funci\u00f3n permisosPorDefecto() contiene listas de permisos hardcodeadas para cada rol. Si la BD tiene permisos, estos se ignoran y se usan los hardcodeados (solo se consulta BD si hay registros en rol_permiso).',
     'C:\\\\backend\\\\src\\\\routes\\\\auth.routes.js:245-273'),
    ('[RED-06] Controladores con estructura casi id\u00e9ntica', 'baja',
     'Todos los controladores sigue el mismo patr\u00f3n try-catch con res.status(500).json(...). Hay mucha repetici\u00f3n de c\u00f3digo.',
     'Todos los controllers en C:\\\\backend\\\\src\\\\controllers'),
]

for title, severity, desc, location in red:
    p = doc.add_paragraph()
    run = p.add_run(f'\u26A0 {title}')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 130, 0)
    doc.add_paragraph(f'[Severidad: {severity.upper()}] {desc}')
    doc.add_paragraph(f'Ubicaci\u00f3n: {location}')

# 6. LO QUE NO FUNCIONA / NO IMPLEMENTADO
doc.add_heading('6. Lo que No Funciona / No Est\u00e1 Implementado', 1)

not_working = [
    ('[NF-01] Router declarativo no implementado', 'alta',
     'app_rtr.dart est\u00e1 vac\u00edo. La navegaci\u00f3n usa Navigator.pushReplacement manual en lugar de GoRouter, Navigator 2.0 o similar.',
     'C:\\\\mobile\\\\lib\\\\core\\\\rtr\\\\app_rtr.dart'),
    ('[NF-02] M\u00f3dulos del Dashboard deshabilitados', 'media',
     'Servicios, Reportes, Operaciones, Estad\u00edsticas y Configuraci\u00f3n tienen enabled: false. No hay implementaci\u00f3n.',
     'C:\\\\mobile\\\\lib\\\\features\\\\dash\\\\dash_scr.dart:51-55'),
    ('[NF-03] Sin gestor de estado global', 'media',
     'No hay Provider, Riverpod, BLoC ni similar. El estado es completamente local con setState.',
     'Toda la app mobile'),
    ('[NF-04] Sin persistencia de sesi\u00f3n (token en memoria vol\u00e1til)', 'media',
     'AuthSession guarda el token en un campo static. Se pierde al recargar/cerrar la app.',
     'C:\\\\mobile\\\\lib\\\\core\\\\auth\\\\auth_session.dart'),
    ('[NF-05] Sin pruebas automatizadas (backend)', 'alta',
     'package.json tiene "test": "echo \\"Error: no test specified\\" && exit 1". No hay un solo test en el backend.',
     'C:\\\\backend\\\\package.json:8'),
    ('[NF-06] Pruebas m\u00ednimas en mobile (1 solo test widget)', 'alta',
     'Solo existe un test widget (\'widget_test.dart\') que probablemente falla por no tener MaterialApp apropiado.',
     'C:\\\\mobile\\\\test\\\\widget_test.dart'),
    ('[NF-07] Sin CI/CD pipeline', 'media',
     'No hay configuraci\u00f3n de GitHub Actions, GitLab CI u otro pipeline de integraci\u00f3n continua.',
     'Ra\u00edz del proyecto'),
    ('[NF-08] Sin Docker ni contenedorizaci\u00f3n', 'baja',
     'No hay Dockerfile ni docker-compose.yml. El despliegue depende del entorno local.',
     'Ra\u00edz del proyecto'),
    ('[NF-09] Sin soporte de archivos m\u00f3vil (stubs retornan null)', 'media',
     'file_pick_stub.dart retorna null. La selecci\u00f3n de archivos no funciona en plataformas no-web.',
     'C:\\\\mobile\\\\lib\\\\core\\\\file\\\\file_pick_stub.dart'),
    ('[NF-10] Sin visualizaci\u00f3n de PDF nativa (solo web)', 'media',
     'pdf_preview_stub.dart no tiene implementaci\u00f3n nativa. Solo funciona en web.',
     'C:\\\\mobile\\\\lib\\\\core\\\\pdf\\\\pdf_preview_stub.dart'),
    ('[NF-11] Bot\u00f3n "Olvid\u00f3 su contrase\u00f1a" sin funcionalidad', 'baja',
     'El bot\u00f3n "Olvid\u00f3 su contrase\u00f1a" en la pantalla de login tiene onPressed: () => {} (vac\u00edo).',
     'C:\\\\mobile\\\\lib\\\\features\\\\auth\\\\auth_scr.dart:252'),
]

for title, severity, desc, location in not_working:
    p = doc.add_paragraph()
    run = p.add_run(f'\u274C {title}')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    doc.add_paragraph(f'[Severidad: {severity.upper()}] {desc}')
    doc.add_paragraph(f'Ubicaci\u00f3n: {location}')

# 7. RECOMENDACIONES DE MEJORA
doc.add_heading('7. Recomendaciones para Mejorar', 1)

improvements = [
    ('7.1 SEGURIDAD (Prioridad: Alta)', [
        'Implementar rate limiting (express-rate-limit) en el endpoint de login',
        'Usar refresh tokens con JWT (access + refresh tokens)',
        'Mover JWT_SECRET a variable de entorno segura y rotarla peri\u00f3dicamente',
        'Agregar validaci\u00f3n de longitud y caracteres en entradas de texto',
        'Implementar bloqueo de cuenta tras N intentos fallidos de login',
        'Sanitizar nombres de tabla en consultas din\u00e1micas (insertarBasico, cambiarActivo)',
        'Nunca exponer el archivo .env con credenciales reales en el repositorio (agregar a .gitignore)',
        'Validar permisos en el endpoint de cambio de contrase\u00f1a'
    ]),
    ('7.2 RENDIMIENTO (Prioridad: Alta)', [
        'Implementar pool de conexiones ODBC en lugar de abrir/cerrar por operaci\u00f3n',
        'Agregar \u00edndices SQL en columnas usadas en WHERE/JOIN (correo_institucional, cedula, evento_personal.evento_id, etc.)',
        'Reducir el l\u00edmite de JSON de 25MB a un valor m\u00e1s razonable (ej. 5-10MB)',
        'Evaluar el uso de Redis para cach\u00e9 de consultas frecuentes (cat\u00e1logos, roles, permisos)'
    ]),
    ('7.3 ARQUITECTURA (Prioridad: Media)', [
        'Migrar a un gestor de estado global en Flutter (Riverpod recomendado por ser moderno y simple)',
        'Implementar GoRouter o Navigator 2.0 para navegaci\u00f3n declarativa con deep linking',
        'Centralizar el manejo de errores HTTP en el ApiClient para redirigir al login en 401',
        'Extraer l\u00f3gica repetida de controladores a un helper/middleware gen\u00e9rico',
        'Unificar las funciones de contrase\u00f1a por defecto en un solo lugar',
        'Eliminar dependencias no utilizadas (msnodesqlv8, mssql) del package.json'
    ]),
    ('7.4 PERSISTENCIA (Prioridad: Media)', [
        'Persistir token JWT en flutter_secure_storage (o al menos SharedPreferences)',
        'Persistir notificaciones le\u00eddas en almacenamiento local (Hive, SQLite, SharedPreferences)',
        'Implementar refresh autom\u00e1tico de token antes de expirar'
    ]),
    ('7.5 CALIDAD Y TESTING (Prioridad: Alta)', [
        'Escribir tests unitarios para services y repositories del backend (usar Jest o Mocha)',
        'Escribir tests unitarios para modelos, servicios y widgets en Flutter',
        'Configurar CI/CD con GitHub Actions (lint, typecheck, test, build)',
        'Configurar linting m\u00e1s estricto tanto en backend (ESLint) como en mobile (flutter_lints ya incluido)'
    ]),
    ('7.6 DOCUMENTACI\u00d3N (Prioridad: Media)', [
        'Completar README.md con instrucciones de instalaci\u00f3n, configuracion y despliegue',
        'Documentar la API REST (OpenAPI/Swagger)',
        'Agregar diagrama de base de datos y arquitectura'
    ]),
    ('7.7 FUNCIONALIDADES FALTANTES (Prioridad: Media)', [
        'Implementar los m\u00f3dulos deshabilitados del Dashboard (Servicios, Reportes, Operaciones, Estad\u00edsticas, Configuraci\u00f3n)',
        'Implementar funcionalidad de "Olvid\u00f3 su contrase\u00f1a" (env\u00edo de correo)',
        'Implementar visualizaci\u00f3n de PDF en plataformas nativas (usar flutter_pdfview u open_file)',
        'Implementar selecci\u00f3n de archivos/im\u00e1genes en plataformas nativas (usar file_picker package)',
        'Agregar soporte para notificaciones push (Firebase Cloud Messaging)'
    ]),
    ('7.8 INFRAESTRUCTURA (Prioridad: Baja)', [
        'Crear Dockerfile y docker-compose.yml para el backend',
        'Configurar entorno de desarrollo con variables de entorno separadas (.env.dev, .env.prod)',
        'Agregar .gitignore adecuado que excluya node_modules, .env, build/', 
        'Crear script de inicializaci\u00f3n r\u00e1pida (setup.sh/setup.ps1)'
    ]),
]

for section_title, items in improvements:
    doc.add_heading(section_title, 2)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run('\u2192 ' + item)
        run.font.color.rgb = RGBColor(0, 70, 180)

# 8. ESTADISTICAS
doc.add_heading('8. Estad\u00edsticas del Proyecto', 1)

stats = [
    ('Archivos totales (sin node_modules)', '~90 archivos'),
    ('Backend (JS)', '27 archivos fuente, ~3,590 l\u00edneas'),
    ('Mobile (Dart/Flutter)', '~40+ archivos fuente, ~5,000+ l\u00edneas'),
    ('Base de Datos (SQL)', '10 scripts de migraci\u00f3n, ~1,511 l\u00edneas'),
    ('Commits en git', '2 commits'),
    ('Dependencias backend', '6 directas (bcrypt, cors, dotenv, express, jsonwebtoken, odbc) + 2 redundantes (msnodesqlv8, mssql)'),
    ('Dependencias mobile', '3 directas (cupertino_icons, http, web)'),
    ('Roles del sistema', '8-9 roles con ~80-100 permisos'),
    ('Insignias (gamificaci\u00f3n)', '15 niveles escalonados'),
    ('Datos de prueba', '58 usuarios de ejemplo'),
]

for label, value in stats:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{label}: ')
    run1.bold = True
    p.add_run(value)

# 9. RESUMEN FINAL
doc.add_heading('9. Resumen Final', 1)

doc.add_paragraph('El proyecto BITSAC (SIGO-GCAM) es un sistema de gesti\u00f3n operativa funcional y bien estructurado '
    'que utiliza tecnolog\u00edas modernas (Node.js/Express 5 + Flutter 3.44). La arquitectura en capas del backend '
    'es s\u00f3lida y el frontend mobile tiene una base multi-plataforma s\u00f3lida con Flutter.')

doc.add_paragraph('Sin embargo, se identificaron los siguientes puntos cr\u00edticos:')

summary_table = doc.add_table(rows=5, cols=2)
summary_table.style = 'Light Grid Accent 1'
summary_table.cell(0,0).text = 'Categor\u00eda'
summary_table.cell(0,1).text = 'Cantidad'
summary_table.cell(1,0).text = 'Errores cr\u00edticos'
summary_table.cell(1,1).text = '5'
summary_table.cell(2,0).text = 'Errores graves/medios'
summary_table.cell(2,1).text = '15+'
summary_table.cell(3,0).text = 'C\u00f3digo obsoleto'
summary_table.cell(3,1).text = '7'
summary_table.cell(4,0).text = 'Redundancias'
summary_table.cell(4,1).text = '6'

doc.add_paragraph('')
doc.add_paragraph('Puntuaci\u00f3n de salud del proyecto: 6.5/10')
doc.add_paragraph('  - Arquitectura: 8/10')
doc.add_paragraph('  - Seguridad: 4/10')
doc.add_paragraph('  - Rendimiento: 5/10')
doc.add_paragraph('  - Testing: 1/10')
doc.add_paragraph('  - Documentaci\u00f3n: 3/10')
doc.add_paragraph('  - Mantenibilidad: 7/10')

doc.add_paragraph('')
doc.add_paragraph('Recomendaci\u00f3n: Abordar primero los 5 errores cr\u00edticos de seguridad (secci\u00f3n 3.1), '
    'luego implementar pool de conexiones y gestor de estado, y finalmente configurar CI/CD con tests automatizados.')

# Save
output_path = 'C:\\Users\\ASUS\\Desktop\\SEGURA EP\\BITSAC\\Reporte_Verificacion_BITSAC.docx'
doc.save(output_path)
print(f'Reporte generado exitosamente: {output_path}')
